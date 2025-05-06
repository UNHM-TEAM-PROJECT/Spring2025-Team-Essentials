import json
import os
import re
import pdfplumber
import zipfile
import tempfile
import hashlib  

from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chat_models import ChatOpenAI
from langchain.schema import Document, HumanMessage
from PIL import Image
import pytesseract
import json
import os
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from io import BytesIO
latest_syllabus_info = {}
buffer = BytesIO()
# -*- coding: utf-8 -*-

# Define the cache for extracted information
extracted_info_cache = {}

SENDER_EMAIL = "Essentials2025UNH@gmail.com"
SENDER_APP_PASSWORD = "prpa flfb znbk oglt"
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
# Final PROMPT_TEMPLATE for a responsive, non-repetitive NECHE Compliance Assistant
PROMPT_TEMPLATE = """
You are a lively, helpful NECHE Syllabus Compliance Assistant, acting like a friendly colleague who’s great at chatting and getting to the point. Your job is to help users check syllabus compliance, extract course/instructor details, and answer all questions clearly and uniquely, with no repetitive phrases. Always sound natural, avoid robotic responses, and keep it concise.

### Core Guidelines:

1. **Natural, Varied Responses**:
   - Use a warm, casual tone (e.g., “Yo! Let’s sort out your syllabus!” or “Hey, I got you!”).
   - For greetings, mix it up:
     - "hi" → “Hey! What’s good?”
     - "hello" → “Yo! Ready to tackle some NECHE stuff?”
     - "hey" → “Hey there! Let’s dive in!”
   - For personal questions, give distinct answers:
     - "how are you" → “I’m pumped to help! You good?”
     - "who are you" → “I’m your syllabus guru, here to nail NECHE compliance and answer your questions!”
     - "tell me about yourself" → “Just a cool AI built to make your syllabus NECHE-ready. I dig into course details and keep things compliant!”
     - "what’s your purpose" → “I’m all about helping your syllabus meet NECHE standards and answering any course questions you’ve got!”

2. **Clear NECHE Explanations**:
   - For “what is NECHE?”: “NECHE is the New England Commission of Higher Education. It sets rules to ensure syllabi have key info like instructor contacts, learning goals, and credit hour details for academic quality.”
   - For “what is NECHE compliance?”: “NECHE compliance means your syllabus has all the required details—like professor’s name, email, office hours, course objectives, and credit policies—to meet accreditation standards.”

3. **Syllabus Compliance**:
   - For “Is this syllabus compliant?”, check the latest syllabus and list missing fields (e.g., “Nope, it’s missing Phone Number and Course SLOs”).
   - For specific details (e.g., “Who’s the professor?”), use exact extracted text (e.g., “It’s Phillip Deen, email: phillip.deen@unh.edu”).
   - If no syllabus is uploaded, say: “Got no syllabus to check yet! Upload one, and I’ll scan it for NECHE compliance.”

4. **General/Casual Questions**:
   - For “what can you do” or “how can you help me”: “I can scan your syllabus for NECHE compliance, grab details like professor info or grading policies, and explain accreditation. Hit me with a syllabus or a question!”
   - For unrelated questions, nudge back: “I’m all about NECHE compliance, but I can chat syllabi or standards. What’s up?”

5. **Response Rules**:
   - Keep it short, clear, and unique—no repeating phrases like “Excited to help” or “I’m your assistant.”
   - Never say “Refer to” or “See the syllabus.”
   - Use “Not Found” for missing syllabus details.
   - Avoid bold or extra formatting unless listing missing fields.
   - If no syllabus exists, prompt: “Upload a syllabus, and I’m on it!”

6. **Supported Topics**:
   - NECHE requirements (instructor details, course objectives, etc.).
   - Syllabus info (professor name, grading, assignments).
   - NECHE and accreditation explanations.
   - Casual education/compliance chats.

7. **Context**:
   - Use the latest uploaded syllabus for answers.
   - Handle follow-ups smoothly (e.g., “What’s their email?” after “Who’s the professor?”).
   - Vary wording to keep it fresh.

### Example Conversation:
- **User**: “Hi”
- **Assistant**: “Hey! What’s good?”
- **User**: “How are you?”
- **Assistant**: “I’m pumped to help! You good?”
- **User**: “Who are you?”
- **Assistant**: “I’m your syllabus guru, here to nail NECHE compliance and answer your questions!”
- **User**: “What is NECHE?”
- **Assistant**: “NECHE is the New England Commission of Higher Education. It sets rules to ensure syllabi have key info like instructor contacts and learning goals.”
- **User**: “Is it compliant?”
- **Assistant**: “Got no syllabus yet! Upload one, and I’ll check it for NECHE compliance.”

### Notes:
- Use latest syllabus data.
- Ensure every answer feels fresh and tailored.
- Stay concise and avoid redundancy.
"""
# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = "uploads"
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Chroma vector store directory
persist_directory = 'db'
os.makedirs(persist_directory, exist_ok=True)

# OpenAI API Key
OPENAI_API_KEY = "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"

llm = ChatOpenAI(
    model="gpt-3.5-turbo",
    openai_api_key=OPENAI_API_KEY,
    temperature=0,
    top_p=1
)

# Load OpenAI Embeddings
embeddings = OpenAIEmbeddings(model="text-embedding-ada-002", openai_api_key=OPENAI_API_KEY)

# Initialize Chroma vector store
db = None

def initialize_chroma():
    global db
    if os.path.exists(os.path.join(persist_directory, "index")):
        db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
    else:
        db = None

initialize_chroma()

# Compliance Criteria: Key Information to Check
required_compliance_items = {
    "Instructor Name": [
        r"(?i)\b(name|instructor|dr\.|mr\.|ms\.|mrs\.)[:\-]?\s*([a-zA-Z]+(?:\s+[a-zA-Z]+)+)(?=,|\s|$)",
        r"(?i)^[a-zA-Z]+(?:\s[a-zA-Z]+)+$"
    ],
    "Title or Rank": [
        r"(?i)\b(professor|assistant professor|associate professor|full professor|senior lecturer|lecturer|instructor|adjunct|faculty|teaching fellow|grad assistant)\b",
        r"(?i)\b(ph\.?d|phd|dr\.)\b",
        r"(?i)\b(professor\s+[a-zA-Z]+(?:\s+[a-zA-Z]+)*,\s*ph\.?d)\b",
        r"(?i)\b([a-zA-Z]+\s+[a-zA-Z]+(?:\s+[a-zA-Z]+)*\s*(professor|ph\.?d))\b",
    ],
    "Preferred Contact Method": [
        r"(?i)\b(contact method|preferred contact|contact information|best way to contact|communication expectations|how to reach)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(email is the best|contact me via|preferred method|reach me by)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Email Address": [
        r"(?i)(email|E-mail|contact)\s*:\s*[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(?=.*(professor|instructor|faculty|contact))",
        r"(?i)(professor|instructor|faculty)\s*[:-]?\s*[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(?=\s*\(?(professor|instructor|faculty|contact)\)?)"
    ],
    "Phone Number": [
        r"(?i)(phone|office information|land line|office phone|contact|cell|direct)\s*[:\-]?\s*\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
        r"\(\d{3}\)\s*\d{3}[-.\s]?\d{4}(?=\s*\(?(office|phone|cell|contact|direct|land line)\)?)",
        r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b(?=.*(professor|phone|cell|instructor|office hours|land line|office information))",
        r"(?i)(phone|office information|land line)[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}"
    ],
    "Office Address": [
        r"(?i)\b(office address|office|room|building)\b\s*[:\-]?\s*(Room\s*\d+[^\n]*|Building\s*[^\n]*|[^\n]+?)(?=\n\n|\Z)",
        r"(?i)\b(office information|contact information)\b\s*[:\-]?\s*(Room\s*\d+[^\n]*|[^\n]+?)(?=\n\n|\Z)"
    ],
    "Office Hours": [
        r"(?i)\b(office hours|hours|availability|by appointment|drop by|open door)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(schedule an appointment|contact me during|available by)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Location (Physical or Remote)": [
        r"(?i)\b(physical location|remote|by appointment|location|online|in person|zoom|in-person|class meetings|room|building)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(classroom|office location|meeting place|class\s+meetings)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(Tuesdays\s+1:10\s*–\s*4:00\s*PM\s+Room\s+503)\b"
    ],
    "Course SLOs": [
        r"(?i)\b(course\s+learning\s+outcomes|student\s+learning\s+outcomes|SLOs|learning\s+objectives|upon completion)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(students\s+will(?:\s+be\s+able\s+to)?)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(course\s+goals|educational\s+objectives)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Credit Hour Workload": [
        r"(?i)\b(workload|credit\s+hour\s+expectations|credit\s+hours|credit\s+hour\s+policy)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(minimum\s+45\s+hours\s+per\s+credit|total\s+workload|time\s+commitment|weekly\s+engagement)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Assignments & Delivery": [
        r"(?i)\b(assignments\s+&\s+delivery|assignment\s+types|graded\s+work|coursework)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(assignments|exams|projects|quizzes|submission)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Grading Procedures & Final Grade Scale": [
        r"(?i)\b(grading\s+procedures\s+&\s+final\s+grade\s+scale|grading\s+policy|grade\s+breakdown|grade\s+scale)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(grading\s+criteria|assessment\s+methods)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Assignment Deadlines & Policies": [
        r"(?i)\b(assignment\s+deadlines\s+&\s+policies|deadlines|late\s+work|submission\s+policies)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(late\s+submissions|due\s+dates|extensions)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Course Number and Title": [
        r"(?i)\b(course\s+number|course\s+title|course\s+name|course\s+code)\b\s*[:\-]?\s*([A-Z]+\s+\d+\s+[^:]+)",
        r"(?i)\b(PSYC\s+502\s+Research\s+Methods\s+in\s+Psychology)\b"
    ],
    "Number of Credits/Units (include a link to the federal definition of a credit hour)": [
        r"(?i)\b(credits|credit\s+hours|number\s+of\s+credits|units|credit\s+hour\s+policy)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(federal\s+definition\s+of\s+a\s+credit\s+hour)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(https://catalog\.unh\.edu/undergraduate/academic-policies-procedures/credit-hour-policy/)\b"
    ],
    "Modality/Meeting Time and Place": [
        r"(?i)\b(modality|meeting\s+time|meeting\s+place|class\s+schedule|class\s+time|location|class\s+meeting)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(Tuesdays\s+1:10\s*–\s*4:00\s*PM\s+Room\s+503)\b"
    ],
    "Semester/Term (and start/end dates)": [
        r"(?i)\b(semester|term|start\s+date|end\s+date|academic\s+term|course\s+schedule)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(Spring\s+2025|1/21\s*–\s*5/7)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(course\s+schedule.*?(1/21\s*–\s*5/7))\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Department/Program": [
        r"(?i)\b(department|program|school|college|division|faculty|institute)\b\s*(of|for)?\s*[a-zA-Z]+",
        r"(?i)\b(Psychology|Neuropsychology)\b"
    ],
    "Format (e.g., lecture plus lab/discussion etc.)": [
        r"(?i)\b(course\s+format|class\s+format|mode\s+of\s+instruction|lecture\s+plus\s+lab|discussion\s+based|seminar|writing\s+intensive|in-class\s+activities)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(writing\s+intensive\s+course|various\s+modalities|in-class\s+activities|research\s+project|exams)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Course Description (minimum course catalog description)": [
        r"(?i)\b(course\s+description|course\s+summary|course\s+overview)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(writing\s+intensive\s+course\s+designed\s+to\s+present\s+information)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Sequence of Course Topics and Important Dates": [
        r"(?i)\b(sequence\s+of\s+course\s+topics\s+and\s+important\s+dates|course\s+schedule|lab\s+schedule|weekly\s+topics|important\s+dates|course\s+timeline)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(module\s+#?\d+|week\s+\d+|jan\s+\d+|feb\s+\d+|mar\s+\d+|apr\s+\d+|may\s+\d+)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(topics\s+and\s+dates|course\s+calendar|schedule\s+of\s+topics)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Required/Recommended Textbook (or other source for course reference information)": [
        r"(?i)\b(required\s+texts?|textbooks?|recommended\s+books?|course\s+readings?|assigned\s+books?)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(Research\s+Methods\s+in\s+Psychology\s+4th\s+American\s+Edition)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)": [
        r"(?i)\b(other\s+required\/recommended\s+materials|required\s+materials|recommended\s+materials|course\s+materials|supplementary\s+materials|software|clicker|remote)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(python\s+\d+\.\d+|virtualbox|docker|laptop\s+with\s+\d+gb\s+ram|antivirus\s+software|specific\s+hardware)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(additional\s+resources|required\s+equipment|recommended\s+tools)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Technical Requirements": [
        r"(?i)\b(technical\s+requirements|reliable\s+internet|recommended\s+browser)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(students\s+must\s+have\s+access\s+to)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Attendance": [
        r"(?i)\b(attendance\s+policy|attendance\s+requirements|absence\s+policy|class\s+attendance)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(regular\s+attendance\s+is\s+expected|absences\s+will)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Academic Integrity/Plagiarism/AI": [
        r"(?i)\b(academic\s+integrity|plagiarism|academic\s+honesty|use\s+of\s+automated\s+writing\s+tools|AI\s+tools)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(Turnitin|unauthorized\s+AI|prohibited\s+AI|ChatGPT|Grammarly|academic\s+dishonesty|UNH\s+academic\s+policy)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(academic\s+integrity/plagiarism/ai)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Course Prerequisites": [
        r"(?i)\b(prerequisites?|pre[-\s]?reqs?)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(PSYC\s+401|PSYC\s+402)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Simultaneous 700/800 Course Designation": [
        r"(?i)\b(simultaneous\s+700/800|dual\s+level\s+course|700/800\s+designation|combined\s+with\s+[A-Z]+\s+(7|8)\d{2})\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "University Requirements": [
        r"(?i)\b(university\s+requirements|academic\s+requirements\s+fulfilled|discovery\s+attribute|writing\s+intensive)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)",
        r"(?i)\b(required\s+course\s+for)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ],
    "Teaching Assistants (Names and Contact Information)": [
        r"(?i)\b(teaching\s+assistants?|TAs?|graduate\s+assistants?|TA\s+contact|assistant\s+name)\b\s*[:\-]?\s*([\s\S]+?)(?=\n\n|\Z)"
    ]
}

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'zip'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

import multiprocessing

def process_page(page):
    page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ''
    if not page_text:
        page_image = page.to_image(resolution=150).original
        page_text = pytesseract.image_to_string(Image.fromarray(page_image), config="--psm 3")
    return page_text

def extract_text_from_pdf(pdf_path):
    extracted_text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ''
                if not page_text:
                    print(f"🔹 No direct text found on page {page.page_number}, using OCR...")
                    page_image = page.to_image(resolution=300).original
                    # Ensure page_image is a PIL Image and pass directly to pytesseract
                    page_text = pytesseract.image_to_string(page_image, config="--psm 6")
                table_text = ""
                if page.extract_tables():
                    for table in page.extract_tables():
                        for row in table:
                            table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                extracted_text.append(page_text.strip() + "\n" + table_text.strip())
    except Exception as e:
        print(f"⚠️ Error processing PDF: {str(e)}")
        return None
    return "\n".join(extracted_text).strip() if extracted_text else None

def extract_text_from_docx(docx_path):
    """
    Extracts structured text from a Word document, ensuring tables are formatted correctly.
    """
    text = []
    try:
        doc = DocxDocument(docx_path)
        
        # ✅ Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())

        # ✅ Extract tables properly
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append(row_text)

    except Exception as e:
        print(f"⚠️ Error processing DOCX: {str(e)}")
        return None

    return "\n".join(text).strip() if text else None

processed_results = {}  # 🔁 GLOBAL dictionary to store all file data

@app.route('/send_email', methods=['POST'])
def send_email():
    try:
        data = request.get_json()
        filename = data.get('filename')
        to_email = data.get('to', "")
        subject = data.get('subject', f'Syllabus Compliance Report - {filename}')
        body = data.get('body', '')

        if not filename or filename not in processed_results:
            return jsonify({"error": "No valid syllabus data to share. Please upload a file first."}), 400

        extracted_info = processed_results[filename]['extracted_information']

        required_fields = [
            "Instructor Name", "Title or Rank", "Preferred Contact Method",
            "Email Address", "Phone Number", "Office Address", "Office Hours", "Location (Physical or Remote)",
            "Course SLOs", "Credit Hour Workload", "Assignments & Delivery", "Grading Procedures & Final Grade Scale",
            "Assignment Deadlines & Policies"
        ]
        minimum_fields = [
            "Course Number and Title", "Number of Credits/Units (include a link to the federal definition of a credit hour)",
            "Modality/Meeting Time and Place", "Semester/Term (and start/end dates)", "Department/Program",
            "Format (e.g., lecture plus lab/discussion etc.)", "Course Description (minimum course catalog description)",
            "Sequence of Course Topics and Important Dates", "Required/Recommended Textbook (or other source for course reference information)",
            "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)", "Technical Requirements",
            "Attendance", "Academic Integrity/Plagiarism/AI"
        ]
        optional_fields = [
            "Course Prerequisites", "Simultaneous 700/800 Course Designation", "University Requirements",
            "Teaching Assistants (Names and Contact Information)"
        ]
        all_fields = required_fields + minimum_fields + optional_fields

        # Course details parsing
        course_title_raw = extracted_info.get("Course Number and Title", "Not Found")
        course_id = ""
        course_name = course_title_raw
        if course_title_raw != "Not Found":
            parts = course_title_raw.split(" ")
            if len(parts) >= 2 and re.match(r'^[A-Z]{2,}\d{2,}', parts[0] + parts[1]):
                course_id = parts[0] + " " + parts[1]
                course_name = " ".join(parts[2:])
            else:
                course_id = course_title_raw
        semester = extracted_info.get("Semester/Term (and start/end dates)", "Not Found")
        if semester != "Not Found":
            semester = semester.split(" ")[0]
        instructor = extracted_info.get("Instructor Name", "Not Found")

        # Generate PDF with ReportLab
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=30, bottomMargin=30, leftMargin=30, rightMargin=30)
        elements = []

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, spaceAfter=10, alignment=1)  # Centered
        detail_style = ParagraphStyle('Detail', parent=styles['Normal'], fontSize=12, spaceAfter=6, alignment=1)  # Centered
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003591')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 14),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#f5f5f5')),
            ('BACKGROUND', (1, 1), (1, -1), colors.white),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ])

        # Heading
        elements.append(Paragraph("Syllabus Compliance Report", title_style))

        # Course details
        course_details = []
        if course_id != "Not Found":
            course_details.append(f"Course ID: {course_id}")
        if course_name != "Not Found":
            course_details.append(f"Course Name: {course_name}")
        if semester != "Not Found":
            course_details.append(f"Semester: {semester}")
        if instructor != "Not Found":
            course_details.append(f"Instructor Name: {instructor}")

        if course_details:
            for detail in course_details:
                elements.append(Paragraph(detail, detail_style))
        else:
            elements.append(Paragraph("No course details available", detail_style))

        # Table
        table_data = [["Syllabus Items", "Syllabus Details"]]
        for field in all_fields:
            value = extracted_info.get(field, "Not Found")
            if field in optional_fields and value == "Not Found":
                continue
            field_label = field
            if field in required_fields:
                field_label = f"<b>{field}</b>"
            elif field in optional_fields:
                field_label = f"<i>{field}</i>"
            value_text = f"<font color=red><b>Not Found</b></font>" if value == "Not Found" else value.replace("\n", "<br />")
            table_data.append([Paragraph(field_label, styles['Normal']), Paragraph(value_text, styles['Normal'])])

        table = Table(table_data, colWidths=[180, 360])
        table.setStyle(table_style)
        elements.append(table)

        doc.build(elements)
        pdf_data = buffer.getvalue()
        buffer.close()

        # Send email
        msg = MIMEMultipart()
        msg['From'] = SENDER_EMAIL
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        attachment = MIMEApplication(pdf_data, _subtype="pdf")
        attachment.add_header('Content-Disposition', 'attachment', filename=f"{filename}_NECHE_Report.pdf")
        msg.attach(attachment)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SENDER_EMAIL, SENDER_APP_PASSWORD)
            server.sendmail(SENDER_EMAIL, to_email, msg.as_string())

        return jsonify({"success": True, "message": f"Email sent with {filename} NECHE report!"})

    except Exception as e:
        return jsonify({"error": f"Failed to send email: {str(e)}"}), 500

def process_uploaded_pdf(file_path, file_name):
    global db
    try:
        extracted_text = extract_text_from_pdf(file_path)
        print(f"Extracted text from {file_name}: {extracted_text[:1000]}...")  # Debug: Print first 100 chars
    except Exception as e:
        print(f"Error extracting text from {file_name}: {str(e)}")
        raise

    # Split into chunks
    chunks = text_splitter.split_text(extracted_text)
    documents = [Document(page_content=chunk, metadata={"source": file_name}) for chunk in chunks]

    # Reset the vector store and add new documents
    if db is not None:
        db.delete_collection()  # Clear existing documents
    db = Chroma.from_documents(documents, embedding=embeddings, persist_directory=persist_directory)
    db.persist()
    print(f"Vector store updated with {len(documents)} documents from {file_name}")  # Debug: Confirm update

# Function to check compliance
def check_neche_compliance(course_info):
    """
    Checks NECHE compliance by verifying required syllabus details and returns exact extracted text.
    """
    required_fields = [
        "Instructor Name", "Title or Rank", "Preferred Contact Method",
        "Email Address", "Phone Number", "Office Address", "Office Hours", 
        "Location (Physical or Remote)", "Course SLOs", "Credit Hour Workload",
        "Assignments & Delivery", "Grading Procedures & Final Grade Scale",
        "Assignment Deadlines & Policies", "Course Number and Title",
        "Number of Credits/Units (include a link to the federal definition of a credit hour)",
        "Modality/Meeting Time and Place", "Semester/Term (and start/end dates)", 
        "Department/Program", "Format (e.g., lecture plus lab/discussion etc.)",
        "Course Description (minimum course catalog description)",
        "Sequence of Course Topics and Important Dates",
        "Required/Recommended Textbook (or other source for course reference information)",
        "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)",
        "Technical Requirements", "Attendance", "Academic Integrity/Plagiarism/AI"
    ]
    optional_fields = [
        "Course Prerequisites", "Simultaneous 700/800 Course Designation", 
        "University Requirements", "Teaching Assistants (Names and Contact Information)"
    ]
    all_fields = required_fields + optional_fields

    # Identify missing required fields
    missing_fields = [field for field in required_fields if course_info.get(field) in ["Not Found", "", None]]

    # Build compliance message
    compliance_message = "NECHE Compliance Check Results:\n\n"
    for field in all_fields:
        value = course_info.get(field, "Not Found")
        status = "Present" if value != "Not Found" else "Not Found"
        compliance_message += f"{field}\t{value}\n"
    
    # Add missing fields summary
    if missing_fields:
        compliance_message += "\nThe syllabus is not compliant. Missing or incomplete required information:\n"
        for field in missing_fields:
            compliance_message += f"- {field}\n"
        if "Academic Integrity/Plagiarism/AI" in missing_fields:
            compliance_message += "\nNote: The 'Academic Integrity/Plagiarism/AI' section was not detected. This may be due to formatting issues or the section being missing in the syllabus. Please ensure it includes terms like 'academic integrity,' 'plagiarism,' or 'AI tools' and is clearly labeled.\n"
    else:
        compliance_message += "\nThe syllabus is fully compliant with all required NECHE information present."

    # Check credit hour link
    credit_hour_field = "Number of Credits/Units (include a link to the federal definition of a credit hour)"
    if credit_hour_field in course_info and course_info[credit_hour_field] != "Not Found":
        if "https://catalog.unh.edu/undergraduate/academic-policies-procedures/credit-hour-policy/" not in course_info[credit_hour_field]:
            if credit_hour_field not in missing_fields:
                missing_fields.append(credit_hour_field)
                compliance_message += "\n⚠️ Missing link to the federal definition of a credit hour. Expected: https://catalog.unh.edu/undergraduate/academic-policies-procedures/credit-hour-policy/\n"

    print(f"🔍 Compliance Check Debug: {compliance_message}")
    return {
        "compliant": not missing_fields,
        "compliance_check": compliance_message,
        "missing_fields": missing_fields
    }

import re

def format_text_as_bullets(text):
    """
    Converts numbered lists (1., 2., 3.) into bullet points.
    Ensures readable formatting.
    """
    formatted_text = []
    sentences = text.split("\n")

    for sentence in sentences:
        sentence = sentence.strip()
        
        if sentence:
            # ✅ Remove numbers (1., 2., 3.) at the beginning
            sentence = re.sub(r"^\d+\.\s*", "• ", sentence)  # Changes "1. Example" → "• Example"
            formatted_text.append(sentence)  
    
    return "\n".join(formatted_text)

import json
from langchain.schema import HumanMessage

def extract_course_information(text):
    global extracted_info_cache
    formatted_text = format_text_as_bullets(text)
    
    # Generate a unique key for the document
    doc_key = hashlib.md5(formatted_text.encode('utf-8')).hexdigest()
    
    # Check cache
    if doc_key in extracted_info_cache:
        print(f"Using cached extraction for document with key {doc_key}")
        return extracted_info_cache[doc_key]
    
    # Pre-extract with regex
    pre_extracted = {}
    for field, patterns in required_compliance_items.items():
        extracted_text = None
        for pattern in patterns:
            matches = re.finditer(pattern, formatted_text, re.MULTILINE | re.DOTALL)
            for match in matches:
                start = match.start()
                # Capture until a major section break
                end = formatted_text.find("\n\n\n", start)
                if end == -1:
                    end = len(formatted_text)
                extracted_text = formatted_text[start:end].strip()
                # Clean up for specific fields
                if field in ["Course SLOs", "Assignments & Delivery", "Grading Procedures & Final Grade Scale", 
                            "Credit Hour Workload", "Academic Integrity/Plagiarism/AI", "Sequence of Course Topics and Important Dates"]:
                    extracted_text = "\n".join(line for line in extracted_text.split("\n") if not line.startswith("• •"))
                # Limit to single line for concise fields
                if field in ["Instructor Name", "Email Address", "Phone Number", "Office Address", "Course Number and Title"]:
                    extracted_text = extracted_text.split("\n")[0].strip()
                pre_extracted[field] = extracted_text
                print(f"🔍 Pre-extracted '{field}': {pre_extracted[field][:100]}...")
                break
        if field not in pre_extracted:
            pre_extracted[field] = "Not Found"
        
        # Normalize specific fields
        if field == "Title or Rank" and pre_extracted.get(field) != "Not Found":
            title_text = pre_extracted[field].lower()
            if any(x in title_text for x in ["professor", "ph.d", "dr."]):
                pre_extracted[field] = title_text.replace("ph.d", "PhD").title()
        if field == "Location (Physical or Remote)" and pre_extracted.get(field) != "Not Found":
            office = pre_extracted.get("Office Address", "Not Found")
            if office != "Not Found" and office not in pre_extracted[field]:
                pre_extracted[field] += f"\nOffice: {office}"
        if field == "Semester/Term (and start/end dates)" and pre_extracted.get(field) == "Not Found":
            schedule_match = re.search(r"(?i)(Spring\s+2025.*?1/21\s*–\s*5/7)", formatted_text, re.DOTALL)
            if schedule_match:
                pre_extracted[field] = "Spring 2025, 1/21–5/7"
        if field == "Academic Integrity/Plagiarism/AI" and pre_extracted.get(field) == "Not Found":
            integrity_match = re.search(r"(?i)(academic\s+integrity.*?(Turnitin|ChatGPT|UNH\s+academic\s+policy))", formatted_text, re.DOTALL)
            if integrity_match:
                pre_extracted[field] = integrity_match.group(0).strip()

    # LLM prompt for refined extraction
    prompt = f"""
You are a strict NECHE syllabus compliance inspector.

Your task is to extract all the following fields from the syllabus content exactly as written, without summarizing, modifying, or adding any phrases that reference other sections, such as "See," "Refer to," "Check," "Consult," "In the syllabus," or similar. For each field:
- Return the exact text if found in the syllabus.
- Return "Not Found" if the field is completely missing or if the text contains any reference to another section.
- Output must be a clean JSON object — no extra comments or text.

Syllabus text:
{formatted_text}

Return JSON with these fields:
{json.dumps({
    "Instructor Name": "",
    "Title or Rank": "",
    "Preferred Contact Method": "",
    "Email Address": "",
    "Phone Number": "",
    "Office Address": "",
    "Office Hours": "",
    "Location (Physical or Remote)": "",
    "Course SLOs": "",
    "Credit Hour Workload": "",
    "Assignments & Delivery": "",
    "Grading Procedures & Final Grade Scale": "",
    "Assignment Deadlines & Policies": "",
    "Course Number and Title": "",
    "Number of Credits/Units (include a link to the federal definition of a credit hour)": "",
    "Modality/Meeting Time and Place": "",
    "Semester/Term (and start/end dates)": "",
    "Department/Program": "",
    "Format (e.g., lecture plus lab/discussion etc.)": "",
    "Course Description (minimum course catalog description)": "",
    "Sequence of Course Topics and Important Dates": "",
    "Required/Recommended Textbook (or other source for course reference information)": "",
    "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)": "",
    "Technical Requirements": "",
    "Attendance": "",
    "Academic Integrity/Plagiarism/AI": "",
    "Course Prerequisites": "",
    "Simultaneous 700/800 Course Designation": "",
    "University Requirements": "",
    "Teaching Assistants (Names and Contact Information)": ""
}, indent=2)}
"""

    response = llm.invoke([HumanMessage(content=prompt)])
    raw_text = response.content.strip()

    # Ensure valid JSON
    if not raw_text.startswith("{"):
        raw_text = raw_text[raw_text.find("{"):]
    if not raw_text.endswith("}"):
        raw_text = raw_text[:raw_text.rfind("}") + 1]
    
    try:
        extracted_info = json.loads(raw_text)
        # Post-process to ensure no vague responses
        all_fields = [
            "Instructor Name", "Title or Rank", "Preferred Contact Method",
            "Email Address", "Phone Number", "Office Address", "Office Hours", 
            "Location (Physical or Remote)", "Course SLOs", "Credit Hour Workload",
            "Assignments & Delivery", "Grading Procedures & Final Grade Scale",
            "Assignment Deadlines & Policies", "Course Number and Title",
            "Number of Credits/Units (include a link to the federal definition of a credit hour)",
            "Modality/Meeting Time and Place", "Semester/Term (and start/end dates)", 
            "Department/Program", "Format (e.g., lecture plus lab/discussion etc.)",
            "Course Description (minimum course catalog description)",
            "Sequence of Course Topics and Important Dates",
            "Required/Recommended Textbook (or other source for course reference information)",
            "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)",
            "Technical Requirements", "Attendance", "Academic Integrity/Plagiarism/AI",
            "Course Prerequisites", "Simultaneous 700/800 Course Designation",
            "University Requirements", "Teaching Assistants (Names and Contact Information)"
        ]
        for field in all_fields:
            if field not in extracted_info or not extracted_info[field] or extracted_info[field].strip() == "":
                extracted_info[field] = pre_extracted.get(field, "Not Found")
            else:
                # Check for vague responses
                value = extracted_info[field].lower()
                vague_phrases = ["refer to", "see the", "check the", "consult the", "in the syllabus", 
                               "see assignments section", "see grading section", "see"]
                if any(phrase in value for phrase in vague_phrases):
                    extracted_info[field] = pre_extracted.get(field, "Not Found")
        
        # Normalize and clean up
        if "Professor's Email Address" in extracted_info:
            extracted_info["Email Address"] = extracted_info.pop("Professor's Email Address")
        if "Professor's Phone Number" in extracted_info:
            extracted_info["Phone Number"] = extracted_info.pop("Professor's Phone Number")
        
        # Ensure Academic Integrity is captured
        if extracted_info["Academic Integrity/Plagiarism/AI"] == "Not Found":
            integrity_match = re.search(r"(?i)(academic\s+integrity.*?(Turnitin|ChatGPT|UNH\s+academic\s+policy))", formatted_text, re.DOTALL)
            if integrity_match:
                extracted_info["Academic Integrity/Plagiarism/AI"] = integrity_match.group(0).strip()
        
        # Cache the extracted information
        extracted_info_cache[doc_key] = extracted_info

    except Exception as e:
        print(f"❌ Failed to parse response: {e}")
        extracted_info = {field: pre_extracted.get(field, "Not Found") for field in all_fields}
        extracted_info_cache[doc_key] = extracted_info

    return extracted_info

from langchain.text_splitter import RecursiveCharacterTextSplitter

# ✅ Define the text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    separators=["\n\n", "\n", " "]
)

@app.route('/upload', methods=['POST'])
def upload_file():
    global latest_syllabus_info, processed_results

    if 'file' not in request.files:
        print("⚠️ No file provided in request")
        return jsonify({"error": "No file provided"}), 400

    files = request.files.getlist('file')
    if not files or all(f.filename == '' for f in files):
        print("⚠️ No valid files selected")
        return jsonify({"error": "No selected file"}), 400

    results = []
    processed_results.clear()

    try:
        required_fields = [
            "Instructor Name", "Title or Rank", 
            "Preferred Contact Method", "Email Address", "Phone Number",
            "Office Address", "Office Hours", "Location (Physical or Remote)",
            "Course SLOs", "Credit Hour Workload", "Assignments & Delivery",
            "Grading Procedures & Final Grade Scale", "Assignment Deadlines & Policies",
            "Course Number and Title", "Number of Credits/Units (include a link to the federal definition of a credit hour)",
            "Modality/Meeting Time and Place", "Semester/Term (and start/end dates)", "Department/Program",
            "Format (e.g., lecture plus lab/discussion etc.)", "Course Description (minimum course catalog description)",
            "Sequence of Course Topics and Important Dates", "Required/Recommended Textbook (or other source for course reference information)",
            "Other Required/Recommended Materials (e.g., software, clicker remote, etc.)", "Technical Requirements",
            "Attendance", "Academic Integrity/Plagiarism/AI"
        ]
        optional_fields = [
            "Course Prerequisites", "Simultaneous 700/800 Course Designation", "University Requirements",
            "Teaching Assistants (Names and Contact Information)"
        ]

        for file in files:
            filename = secure_filename(file.filename)
            if not allowed_file(filename):
                print(f"⚠️ Invalid file type for {filename}")
                continue

            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            print(f"📄 Saved file: {filename}")

            if filename.endswith('.zip'):
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    with tempfile.TemporaryDirectory() as temp_dir:
                        zip_ref.extractall(temp_dir)
                        for root, _, zip_files in os.walk(temp_dir):
                            for name in zip_files:
                                if name.lower().endswith(('.pdf', '.docx')):
                                    inner_path = os.path.join(root, name)
                                    text = extract_text_from_pdf(inner_path) if name.endswith('.pdf') else extract_text_from_docx(inner_path)
                                    if not text:
                                        print(f"⚠️ No text extracted from {name}")
                                        continue
                                    print(f"📝 Extracted text from {name}: {text[:100]}...")
                                    extracted_info = extract_course_information(text)
                                    # Additional check for vague phrases
                                    vague_phrases = ["refer to", "see the", "check the", "consult the", "in the syllabus", "see assignments section", "see grading section", "see"]
                                    for field in required_fields + optional_fields:
                                        value = extracted_info.get(field, "").strip()
                                        if not value or value.lower() in ["n/a", "na", "not applicable", "none"] or any(phrase in value.lower() for phrase in vague_phrases):
                                            extracted_info[field] = "Not Found"
                                    for key, value in extracted_info.items():
                                        extracted_info[key] = str(value) if not isinstance(value, str) else value
                                    compliance = check_neche_compliance(extracted_info)
                                    processed_results[name] = {
                                        "extracted_information": extracted_info,
                                        "compliance_check": compliance["compliance_check"],
                                        "missing_fields": compliance["missing_fields"]
                                    }
                                    results.append({
                                        "filename": name,
                                        "extracted_information": extracted_info,
                                        "compliance_check": compliance["compliance_check"],
                                        "missing_fields": compliance["missing_fields"]
                                    })
                os.remove(file_path)
                continue

            if filename.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                extracted_text = extract_text_from_docx(file_path)
            else:
                print(f"⚠️ Unsupported file type for {filename}")
                os.remove(file_path)
                continue

            if not extracted_text:
                print(f"⚠️ No text extracted from {filename}")
                os.remove(file_path)
                continue

            print(f"📝 Extracted text from {filename}: {extracted_text[:100]}...")
            extracted_info = extract_course_information(extracted_text)
            # Additional check for vague phrases
            vague_phrases = ["refer to", "see the", "check the", "consult the", "in the syllabus", "see assignments section", "see grading section", "see"]
            for field in required_fields + optional_fields:
                value = extracted_info.get(field, "").strip()
                if not value or value.lower() in ["n/a", "na", "not applicable", "none"] or any(phrase in value.lower() for phrase in vague_phrases):
                    extracted_info[field] = "Not Found"
            for key, value in extracted_info.items():
                extracted_info[key] = str(value) if not isinstance(value, str) else value
            compliance_check_result = check_neche_compliance(extracted_info)
            latest_syllabus_info.clear()
            latest_syllabus_info.update(extracted_info)
            processed_results[filename] = {
                "extracted_information": extracted_info,
                "compliance_check": compliance_check_result["compliance_check"],
                "missing_fields": compliance_check_result["missing_fields"]
            }
            results.append({
                "filename": filename,
                "extracted_information": extracted_info,
                "compliance_check": compliance_check_result["compliance_check"],
                "missing_fields": compliance_check_result["missing_fields"]
            })
            os.remove(file_path)

        if not results:
            print("⚠️ No valid files processed")
            return jsonify({"error": "No valid files processed inside the upload"}), 400

        return jsonify({
            "success": True,
            "message": f"Processed {len(results)} file(s) successfully",
            "results": results
        })

    except Exception as e:
        print(f"❌ Error processing upload: {str(e)}")
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 500

@app.route('/ask', methods=['POST'])
def ask():
    global latest_syllabus_info
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({"error": "No message provided"}), 400

        message = data.get('message', '').strip().lower()

        # Casual responses for greetings and small talk
        casual_responses = {
            "hey": "Hey! How’s your day going?",
            "hello": "Hello! Hope you're having a great day.",
            "hi": "Hi there! How can I assist you today?",
            "how are you": "I'm doing great! Thanks for asking. How about you?",
            "what's up": "Not much, just here to assist with NECHE compliance! What’s up with you?",
            "who are you": "I'm your assistant for NECHE syllabus compliance. I help check syllabus requirements and guide you on accreditation standards.",
            "what do you do": "I assist with NECHE compliance by checking syllabi for required information. Let me know if you need help with that!"
        }

        # Handle casual greetings
        if message in casual_responses:
            return jsonify({"response": casual_responses[message]})

        # Handle professor-related questions
        if "who is the professor" in message or "professor's name" in message:
            professor_name = latest_syllabus_info.get("Instructor Name", "Not Found")
            if professor_name == "Not Found":
                return jsonify({"response": "No professor information found in the most recent syllabus. Please upload a valid syllabus."})
            return jsonify({"response": f"The professor listed in the most recent syllabus is {professor_name}."})

        # Handle compliance check questions
        if "is this syllabus compliant" in message:
            compliance_result = check_neche_compliance(latest_syllabus_info)
            return jsonify({"response": compliance_result["compliance_check"]})

        # NECHE compliance-related keywords
        neche_keywords = [
            "neche", "syllabus", "compliance", "instructor", "credit hours",
            "grading policy", "program SLOs", "assignments", "office hours",
            "submission", "deadlines", "policies", "assessment", "course objectives"
        ]

        # General inquiries about bot's purpose
        general_inquiries = [
            "how can you assist me", "what can you do for me", "tell me about yourself", "what is your purpose"
        ]

        is_neche_related = any(keyword in message for keyword in neche_keywords) or any(inquiry in message for inquiry in general_inquiries)

        # Construct the prompt for other NECHE-related questions
        if is_neche_related or latest_syllabus_info:
            syllabus_summary = "\n".join([f"{k}: {v}" for k, v in latest_syllabus_info.items()])
            prompt = PROMPT_TEMPLATE + f"""
            Latest NECHE Compliance Information (from uploaded syllabus):
            {json.dumps(latest_syllabus_info, indent=2)}

            User's Question: {message}

            Response Guidelines:
            - Provide the exact extracted text from the syllabus for the requested field.
            - Do not use phrases like "Refer to the syllabus" or "See the syllabus."
            - If the information is missing, state "Not Found" for that field.
            - Keep responses short, professional, and NECHE-focused.
            """
        else:
            prompt = """
            I specialize in NECHE syllabus compliance.
            I can check if your syllabus meets NECHE standards, identify missing elements, 
            and guide you on compliance improvements.

            Please ask about syllabus requirements, instructor details, grading policies, or coursework submissions.
            """

        # Invoke LLM
        response = llm.invoke([HumanMessage(content=prompt)])
        return jsonify({"response": response.content.strip()})

    except Exception as e:
        print(f"Error processing question: {str(e)}")
        return jsonify({"error": f"Failed to process question: {str(e)}"}), 500
    
from flask import send_file, Flask, request, jsonify, render_template

@app.route("/download_all_reports_zip", methods=["GET"])
def download_all_reports_zip():
    zip_path = "/path/to/generated/NECHE_All_Reports.zip"  # Update this path accordingly
    return send_file(zip_path, as_attachment=True)

# Serve frontend page
@app.route('/', methods=['GET','POST'])
def home():
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8001)